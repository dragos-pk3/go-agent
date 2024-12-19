from docx import Document
import json

class DocxReplace:
    def __init__(self):
        self.user_config = self.read_json("_internal/__userfiles__/user_config.json")
        self.template_path = self.user_config["CONTRACT_TEMPLATE_PATH"]
        self.user_preferences = self.read_json("_internal/__userfiles__/user_preferences.json")


    def read_json(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        return data

    def read_first_line_as_str(self, filename):
        with open(filename, 'r') as file:
            first_line = file.readline().strip()
        return first_line

    def overwrite_file_with_number(self, filename, number):
        with open(filename, 'w') as file:
            file.write(str(number))

    def load_template(self, file_path):
        return Document(file_path)

    def replace_text_in_run(self, run, data):
        for key, value in data.items():
            run.text = run.text.replace(key, value)

    def create_document(self,doc, data):
        for para in doc.paragraphs:
            for run in para.runs:
                self.replace_text_in_run(run, data)

    def add_image_to_document(self, paragraph, img_path):
        run = paragraph.add_run()
        run.add_picture(img_path)

    def save_document(self, doc, folderPath, filename):
        doc.save(f"{folderPath}/{filename}")

    def replace_slash_with_dash(self, input_string):
        return input_string.replace('/', '-')

    def process_template(self, docData, isContractPJ):
        self.user_config = self.read_json("_internal/__userfiles__/user_config.json")
        self.user_preferences = self.read_json("_internal/__userfiles__/user_preferences.json")
        if isContractPJ:
            self.template_path = self.user_config["CONTRACT_PJ_TEMPLATE_PATH"]
        elif not isContractPJ:
            self.template_path = self.user_config["CONTRACT_TEMPLATE_PATH"]
        else:
            raise ValueError("Invalid contract type")
        
        template = self.load_template(self.template_path)
        self.create_document(template, docData)
        self.add_image_to_document(template.paragraphs[-4], self.user_config["SIGNATURE_PATH"])
        if isContractPJ:
            filename = self.replace_slash_with_dash(f"{docData['NrDoc']} Contract {docData['numeFirma']} ({docData['numeClient']} {docData['prenumClient']}).docx")
        else:
            filename = self.replace_slash_with_dash(f"{docData['NrDoc']} Contract {docData['numeClient']} {docData['prenumClient']}.docx")

        self.save_document(template, self.user_config["OUTPUT_PATH"], filename)
        factura_file_path = "_internal/__userfiles__/factura.txt"
        new_line = f"c/v conf. contract inchiriere nr. {docData['NrDoc']}/{docData['todayDate']} - inchiriere autorulota {docData['startDate'][:5]} - {docData['endDate']}"

        with open(factura_file_path, 'r+') as file:
            content = file.read()
            file.seek(0, 0)
            file.write(new_line.rstrip('\r\n') + '\r\n' + content)
